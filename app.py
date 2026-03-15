import streamlit as st
import pandas as pd
from janome.tokenizer import Tokenizer
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import io
import base64
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt, Mm
import re

# --- 1. ページ設定 ---
st.set_page_config(page_title="Reflection Analyzer", layout="wide")
st.title("📊 振り返り達成度分析システム")

# --- 2. API設定 ---
try:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception:
    st.error("システム設定を確認してください。")
    st.stop()

# --- 3. サイドバー設定 ---
st.sidebar.header("📋 行事・学習の設定")
# 入力欄の初期値に(例)を含める
event_name = st.sidebar.text_input("行事名・授業名", value="(例)理科校外学習")
event_date = st.sidebar.date_input("実施日")

st.sidebar.subheader("学習のねらい")
# ねらいの入力欄にも(例)を含めた初期値を設定
g1 = st.sidebar.text_area("ねらい1", value="(例)理科への興味・関心を高める。")
g2 = st.sidebar.text_area("ねらい2", value="(例)進路意識を高める。")
g3 = st.sidebar.text_area("ねらい3", value="")
g4 = st.sidebar.text_area("ねらい4", value="")

goals = [g for g in [g1, g2, g3, g4] if g.strip()]

t = Tokenizer()

# --- 4. 単語抽出エンジン ---
def extract_words(text):
    words = []
    target_feelings = ["面白い", "楽しい", "凄い", "すごい", "わかる", "驚く", "難しい", "疲れる", "つまらない", "嫌だ", "迷う"]
    if not text or pd.isna(text): return ""
    tokens = list(t.tokenize(str(text)))
    i = 0
    while i < len(tokens):
        token = tokens[i]
        part = token.part_of_speech.split(",")[0]
        base = token.base_form
        surface = token.surface
        if part == "名詞" and len(base) >= 2:
            if base not in ["こと", "もの", "よう", "そう", "これ", "それ"]: words.append(base)
        elif base in target_feelings or part == "形容詞":
            word_to_add = base
            if i + 1 < len(tokens):
                next_t = tokens[i+1]
                if next_t.base_form in ["ない", "ぬ", "ん"] or "助動詞" in next_t.part_of_speech:
                    if next_t.base_form in ["ない", "ぬ", "ん"]:
                        word_to_add = base[:-1] + "くない" if part == "形容詞" else base + "ない"
                        i += 1
            if word_to_add not in ["ある", "する", "いる"]: words.append(word_to_add)
        elif surface == "行き" and i + 1 < len(tokens) and "たい" in tokens[i+1].surface:
            word_to_add = "行きたくない" if i + 2 < len(tokens) and tokens[i+2].base_form == "ない" else "行きたい"
            words.append(word_to_add)
            i += 1
        i += 1
    return " ".join(words)

# --- 5. Word作成用関数 (図の最大化・余白2cm設定) ---
def create_word(evaluation_text, img_bytes, event, date, goal_list):
    doc = Document()
    
    # 余白設定 (上下左右 20mm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
    
    # タイトル
    doc.add_heading(f'{event} 実施報告', 0)
    doc.add_paragraph(f'実施日：{date.strftime("%Y年%m月%d日")}')
    
    # 1. 学習のねらい
    doc.add_heading('1. 学習のねらい', level=1)
    for i, g in enumerate(goal_list):
        p = doc.add_paragraph()
        p.add_run(f"({i+1}) {g}")
    
    # 2. 生徒の反応
    doc.add_heading('2. 生徒の反応（語彙分析）', level=1)
    img_stream = io.BytesIO(img_bytes)
    # 幅を最大化 (6.5インチ)
    doc.add_picture(img_stream, width=Inches(6.5))
    p_note = doc.add_paragraph()
    p_note.add_run("※生徒感想の原文は別紙参照").italic = True
    
    # 3. 評価と考察
    doc.add_heading('3. ねらいに対する評価と考察', level=1)
    
    # テキスト整形
    clean_text = re.sub(r'#+\s*', '', evaluation_text)
    clean_text = clean_text.replace('**', '')
    clean_text = re.sub(r'【報告書】|【行事名】:.*|【設定されたねらい】:|- ねらい\d:.*', '', clean_text).strip()
    
    lines = clean_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
            
        p = doc.add_paragraph()
        # 項目見出しのみを太字にする
        is_heading = any(line.startswith(key) for key in ["【分析】", "【今後の課題・改善案】", "【まとめ】", "ねらい"])
        
        if is_heading:
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
        else:
            p.add_run(line)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- 6. メイン処理 ---
uploaded_file = st.file_uploader("ファイルをアップロードしてください", type=["xlsx", "csv"])

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
    target_col = st.selectbox("分析対象の列を選択してください", df.columns)
    
    if st.button("報告書原稿を作成する"):
        with st.spinner("データを分析中..."):
            all_text_list = df[target_col].dropna().astype(str).tolist()
            all_text_full = "\n".join(all_text_list)
            wakati = extract_words(all_text_full)
            
            if not wakati.strip():
                st.warning("有効な単語が抽出できませんでした。")
            else:
                FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
                wc = WordCloud(font_path=FONT_PATH, background_color="white", width=1200, height=600).generate(wakati)
                img_buf = io.BytesIO()
                wc.to_image().save(img_buf, format='PNG')
                img_bytes = img_buf.getvalue()
                
                st.subheader("📊 振り返り語彙の分析")
                st.image(img_bytes)

                base64_image = base64.b64encode(img_bytes).decode('utf-8')
                context_text = all_text_full[:2000]
                goals_str = "\n".join([f"ねらい{i+1}: {g}" for i, g in enumerate(goals)])

                prompt = f"""
                あなたは学校の教諭として，生徒の振り返り結果の分析報告書を作成してください。
                
                【分析対象行事】: {event_name}
                【設定したねらい】:
                {goals_str}
                
                【生徒の記述内容（抜粋）】:
                {context_text}

                ### 記述のルール
                1. 冒頭に行事名やねらいを再度書く必要はありません。直接「【分析】」から書き始めてください。
                2. 語尾は「〜である」「〜した」等の公的文書の体裁にしてください。
                3. 各ねらいに対して「【分析】」と「【今後の課題・改善案】」を明確に分けて記述してください。
                4. 最後に全体を通じた「【まとめ】」を添えてください。
                5. 記号（#や*）は使用しないでください。
                """

                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}}]}],
                )
                
                evaluation_text = response.choices[0].message.content
                st.divider()
                st.header("📝 達成度評価と考察")
                st.markdown(evaluation_text)
                
                word_file = create_word(evaluation_text, img_bytes, event_name, event_date, goals)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("画像を保存", data=img_bytes, file_name=f"{event_name}_分析.png")
                with col2:
                    st.download_button(
                        label="📄 報告書(Word)を保存",
                        data=word_file,
                        file_name=f"{event_name}_実施報告書.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
